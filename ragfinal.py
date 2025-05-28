import os
import streamlit as st
import tempfile
from typing_extensions import TypedDict, List
import io
import fitz  # PyMuPDF for PDF handling
import pandas as pd
import csv

try:
    import docx
except ImportError:
    # Make docx optional
    docx = None

from langchain_community.chat_models import ChatOpenAI
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_core.vectorstores import InMemoryVectorStore
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import START, StateGraph

# Set page configuration
st.set_page_config(
    page_title="Hytopia Insights",
    page_icon="🔍",
    layout="wide",
)

# App title and description
st.title("💬 Specific Insights")
st.markdown("Upload documents and ask questions to get AI-powered answers!")

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []

if "vector_store" not in st.session_state:
    st.session_state.vector_store = None

if "graph" not in st.session_state:
    st.session_state.graph = None

if "documents" not in st.session_state:
    st.session_state.documents = []

if "file_processor_ready" not in st.session_state:
    st.session_state.file_processor_ready = False


# Helper functions for document processing
def extract_text_from_pdf(file_content):
    """Extract text from PDF file content."""
    text = ""
    with fitz.open(stream=file_content, filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text


def extract_text_from_docx(file_content):
    """Extract text from DOCX file content."""
    if docx is None:
        st.warning("python-docx is not installed. DOCX processing is disabled.")
        return "DOCX processing is disabled. Please install python-docx library."

    doc = docx.Document(io.BytesIO(file_content))
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text


def extract_text_from_csv(file_content):
    """Extract text from CSV file content."""
    text = ""
    csv_data = pd.read_csv(io.BytesIO(file_content))
    # Convert the DataFrame to a string representation
    text = csv_data.to_string(index=False)
    return text


def extract_text_from_txt(file_content):
    """Extract text from TXT file content."""
    return file_content.decode('utf-8')


def process_file(uploaded_file):
    """Process an uploaded file and extract its text content."""
    file_extension = uploaded_file.name.split('.')[-1].lower()
    file_content = uploaded_file.getvalue()

    try:
        if file_extension == 'pdf':
            text = extract_text_from_pdf(file_content)
        elif file_extension == 'docx':
            text = extract_text_from_docx(file_content)
        elif file_extension == 'csv':
            text = extract_text_from_csv(file_content)
        elif file_extension == 'txt':
            text = extract_text_from_txt(file_content)
        else:
            # Try to read as text for other file types
            text = file_content.decode('utf-8')

        return Document(page_content=text, metadata={"source": uploaded_file.name})
    except Exception as e:
        st.error(f"Error processing file {uploaded_file.name}: {str(e)}")
        return None


# Sidebar for API key and file uploads
with st.sidebar:
    st.header("Setup")
    openai_api_key = st.text_input("OpenAI API Key", type="password")

    # File uploader
    st.header("Document Upload")
    uploaded_files = st.file_uploader(
        "Upload documents (PDF, DOCX, CSV, TXT)",
        accept_multiple_files=True,
        type=["pdf", "docx", "csv", "txt"]
    )

    # Process files button
    if st.button("Process Files"):
        if not openai_api_key:
            st.error("Please enter your OpenAI API key!")
        elif not uploaded_files:
            st.error("Please upload at least one file!")
        else:
            with st.spinner("Processing files..."):
                try:
                    # Process each uploaded file
                    processed_documents = []
                    for uploaded_file in uploaded_files:
                        doc = process_file(uploaded_file)
                        if doc:
                            processed_documents.append(doc)

                    if processed_documents:
                        st.session_state.documents = processed_documents
                        st.session_state.file_processor_ready = True
                        st.success(f"Successfully processed {len(processed_documents)} documents!")
                    else:
                        st.error("No documents were successfully processed.")
                except Exception as e:
                    st.error(f"Error processing files: {str(e)}")

    # Initialize RAG system button
    if st.button("Initialize RAG System"):
        if not openai_api_key:
            st.error("Please enter your OpenAI API key!")
        elif not st.session_state.file_processor_ready:
            st.error("Please process files first!")
        else:
            with st.spinner("Initializing RAG system..."):
                try:
                    # Set the API key
                    os.environ["OPENAI_API_KEY"] = openai_api_key

                    # Initialize components
                    llm = ChatOpenAI(model="gpt-4o-mini")
                    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
                    vector_store = InMemoryVectorStore(embeddings)

                    # Split documents into chunks
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=500,
                        chunk_overlap=100,
                        add_start_index=True,
                    )
                    splits = text_splitter.split_documents(st.session_state.documents)

                    # Index documents
                    vector_store.add_documents(documents=splits)

                    # Create a custom RAG prompt template instead of using hub
                    system_template = """
You are a content and communications assistant for Denarii Labs. You have access to Denarii Labs’ internal meeting notes, client info, and social archives, and your job is to write tweets, LinkedIn posts, reports, and strategic content that reflects the Denarii Labs voice and values.

Your writing must strictly follow the tone, formatting, and topic preferences outlined below. If you use general knowledge or outside context, **you must clearly state that it's external**.

---

🧵 **TWITTER GUIDELINES**

**Tone & Voice**
- Curious, never conclusive
- Sharp and confident, but not smug
- Conversational and intelligent
- Use “builder” tone — write like a founder talking to founders
- Assume intelligence in reader, avoid spoon-feeding

**Style Rules**
- Short, punchy, high-signal lines (1–3 per tweet)
- Use line breaks for rhythm and clarity
- Emojis: sparse but strategic (🧵, 🤔, 💡, 👀, etc.)
- Sentence fragments okay. No fluff. No filler.
- Avoid buzzwords unless you’re flipping them on their head

**Proven Formats**
- Tension/contrast: “Most founders do X. Great ones do Y.”
- Insight drops: “Your CAC isn’t broken. Your narrative is.”
- Counterintuitive takes: “The best fundraising advice? Stop optimizing for VCs.”
- Prompt tweets: “What startup idea is stuck in your head lately?”
- Mini-threads: Hook → 3-5 insights → CTA or punch

**Topics We Care About**
- Founder psychology, early-stage GTM, product-market fit
- Infrastructure and tools for builders
- Fundraising, token models, GP/LP dynamics
- Bold startup advice that actually applies at 0→1

**Avoid**
- Corporate tone, “announcement voice”
- Thread clichés (“Here’s the playbook,” “Steal this framework”)
- Empty hype or fluff
- Excessive emojis or hashtag use

---

🔗 **LINKEDIN GUIDELINES**

**Tone & Voice**
- Builder-forward, community-oriented
- Grounded but optimistic
- Specific, sincere, and occasionally personal
- Ambitious but never performative

**Post Structure**
- Start with the *why* or a clear takeaway in line 1
- 2–5 short paragraphs max
- Use line breaks to space thoughts
- Bullet points and emojis okay (🔹, 🙌, 🚀) — keep clean
- Links should be framed with context, not just dumped

**Common Formats**
- Launch or milestone announcements: “We just launched X...”
- Friday roundup with community & partner shoutouts
- Candid reflections on founder lessons or experiments
- Smart hiring posts that reflect mission and purpose
- Event announcements and speaker callouts

**Topics We Focus On**
- What we’re building at Denarii Labs and why
- Macro startup trends, token ecosystems, and early-stage innovation
- Collaborations with Red Beard Ventures and other partners
- Product, ops, and GTM insights that others can learn from

**Avoid**
- Reposting tweets directly without rewriting for LinkedIn
- Overly polished or sales-y tone
- Vague updates with no story or value

**Recurring Phrases/Voice Tics (Use Sparingly)**
- “We’re building in public...”
- “This one was fun.”
- “Tag a founder who needs this.”
- “Not investment advice... but”

---

📌 Reminder: Whenever you use general knowledge, pretraining data, or web-scale info not explicitly provided in Denarii Labs data, indicate it clearly in your response.

                    Context:
                    {context}

                    Question:
                    {question}

                    Answer:
                    """

                    human_template = "{question}"

                    chat_prompt = ChatPromptTemplate.from_messages([
                        ("system", system_template),
                        ("human", human_template)
                    ])


                    # Define application state
                    class State(TypedDict):
                        question: str
                        context: List[Document]
                        answer: str


                    # Define application steps
                    def retrieve(state: State):
                        retrieved_docs = vector_store.similarity_search(state["question"], k=4)
                        return {"context": retrieved_docs}


                    def generate(state: State):
                        docs_content = "\n\n".join(doc.page_content for doc in state["context"])
                        # Create messages directly
                        messages = [
                            {"type": "system",
                             "content": f""""
You are a content and communications assistant for Denarii Labs. You have access to Denarii Labs’ internal meeting notes, client info, and social archives, and your job is to write tweets, LinkedIn posts, reports, and strategic content that reflects the Denarii Labs voice and values.

Your writing must strictly follow the tone, formatting, and topic preferences outlined below. If you use general knowledge or outside context, **you must clearly state that it's external**.

---

🧵 **TWITTER GUIDELINES**

**Tone & Voice**
- Curious, never conclusive
- Sharp and confident, but not smug
- Conversational and intelligent
- Use “builder” tone — write like a founder talking to founders
- Assume intelligence in reader, avoid spoon-feeding

**Style Rules**
- Short, punchy, high-signal lines (1–3 per tweet)
- Use line breaks for rhythm and clarity
- Emojis: sparse but strategic (🧵, 🤔, 💡, 👀, etc.)
- Sentence fragments okay. No fluff. No filler.
- Avoid buzzwords unless you’re flipping them on their head

**Proven Formats**
- Tension/contrast: “Most founders do X. Great ones do Y.”
- Insight drops: “Your CAC isn’t broken. Your narrative is.”
- Counterintuitive takes: “The best fundraising advice? Stop optimizing for VCs.”
- Prompt tweets: “What startup idea is stuck in your head lately?”
- Mini-threads: Hook → 3-5 insights → CTA or punch

**Topics We Care About**
- Founder psychology, early-stage GTM, product-market fit
- Infrastructure and tools for builders
- Fundraising, token models, GP/LP dynamics
- Bold startup advice that actually applies at 0→1

**Avoid**
- Corporate tone, “announcement voice”
- Thread clichés (“Here’s the playbook,” “Steal this framework”)
- Empty hype or fluff
- Excessive emojis or hashtag use

---

🔗 **LINKEDIN GUIDELINES**

**Tone & Voice**
- Builder-forward, community-oriented
- Grounded but optimistic
- Specific, sincere, and occasionally personal
- Ambitious but never performative

**Post Structure**
- Start with the *why* or a clear takeaway in line 1
- 2–5 short paragraphs max
- Use line breaks to space thoughts
- Bullet points and emojis okay (🔹, 🙌, 🚀) — keep clean
- Links should be framed with context, not just dumped

**Common Formats**
- Launch or milestone announcements: “We just launched X...”
- Friday roundup with community & partner shoutouts
- Candid reflections on founder lessons or experiments
- Smart hiring posts that reflect mission and purpose
- Event announcements and speaker callouts

**Topics We Focus On**
- What we’re building at Denarii Labs and why
- Macro startup trends, token ecosystems, and early-stage innovation
- Collaborations with Red Beard Ventures and other partners
- Product, ops, and GTM insights that others can learn from

**Avoid**
- Reposting tweets directly without rewriting for LinkedIn
- Overly polished or sales-y tone
- Vague updates with no story or value

**Recurring Phrases/Voice Tics (Use Sparingly)**
- “We’re building in public...”
- “This one was fun.”
- “Tag a founder who needs this.”
- “Not investment advice... but”

---

📌 Reminder: Whenever you use general knowledge, pretraining data, or web-scale info not explicitly provided in Denarii Labs data, indicate it clearly in your response...

Context:
{context}

Question:
{question}

Answer:
""".strip()

    messages = [
        {"type": "system", "content": system_content.format(context=docs_content, question=state['question'])},
        {"type": "human", "content": state["question"]}
    ]
    response = llm.invoke(messages)
    return {"answer": response.content}

                    # Build the graph
                    graph_builder = StateGraph(State).add_sequence([retrieve, generate])
                    graph_builder.add_edge(START, "retrieve")
                    graph = graph_builder.compile()

                    # Store in session state
                    st.session_state.vector_store = vector_store
                    st.session_state.graph = graph

                    st.success("RAG system initialized successfully! You can now ask questions.")

                except Exception as e:
                    st.error(f"Error initializing RAG system: {str(e)}")

    st.markdown("### About")
    st.markdown(
        "This application uses Retrieval Augmented Generation (RAG) "
        "to answer questions about your uploaded documents. "
        "It retrieves relevant information from the documents and "
        "generates accurate answers based on the content."
    )

# Display uploaded documents section
if st.session_state.documents:
    with st.expander("Uploaded Documents"):
        for i, doc in enumerate(st.session_state.documents):
            st.write(f"{i + 1}. {doc.metadata.get('source', 'Unknown document')}")

# Display chat messages
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])
        if "sources" in message and message["sources"]:
            with st.expander("View Sources"):
                for i, source in enumerate(message["sources"]):
                    st.markdown(f"**Source {i + 1}:**")
                    st.markdown(source)

# Chat input and processing
if prompt := st.chat_input("Ask a question about your documents"):
    # Add user message to chat history
    st.session_state.messages.append({"role": "user", "content": prompt})

    # Display user message
    with st.chat_message("user"):
        st.markdown(prompt)

    # Check if system is initialized
    if st.session_state.graph is None:
        with st.chat_message("assistant"):
            system_message = "⚠️ Please upload documents and initialize the RAG system first using the buttons in the sidebar."
            st.markdown(system_message)
            st.session_state.messages.append({"role": "assistant", "content": system_message})
    else:
        # Display assistant response
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                # Invoke the RAG pipeline
                result = st.session_state.graph.invoke({"question": prompt})
                answer = result["answer"]
                context = result["context"]

                # Display the answer
                st.markdown(answer)

                # Prepare sources for expander
                sources = []
                for i, doc in enumerate(context):
                    source_doc = doc.metadata.get('source', 'Unknown source')
                    start_index = doc.metadata.get('start_index', 'N/A')
                    source_text = f"*From document: {source_doc} (index: {start_index})*\n\n{doc.page_content}"
                    sources.append(source_text)

                # Show sources in expander
                if sources:
                    with st.expander("View Sources"):
                        for i, source in enumerate(sources):
                            st.markdown(f"**Source {i + 1}:**")
                            st.markdown(source)

                # Add to chat history
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": answer,
                    "sources": sources
                })

# Instructions at the bottom
with st.expander("How to use this app"):
    st.markdown("""
    1. Enter your OpenAI API key in the sidebar
    2. Upload your documents (PDF, DOCX, CSV, TXT files)
    3. Click the "Process Files" button
    4. Click the "Initialize RAG System" button
    5. Ask questions about your documents in the chat input
    6. View AI-generated answers and their sources

    **Example questions you can ask:**
    - What are the main topics covered in these documents?
    - Can you summarize the key points about [specific topic]?
    - What does the document say about [specific term or concept]?
    - How are [topic A] and [topic B] related according to the documents?
    - What evidence is provided for [specific claim]?
    """)
